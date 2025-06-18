import os
import time
import mimetypes
from PIL import Image, UnidentifiedImageError
from PIL.ExifTags import TAGS
import exifread
import PyPDF2
import docx
import openpyxl
import magic
import re

# Import new libraries for audio and video metadata extraction
try:
    import hachoir.parser
    import hachoir.metadata
    from hachoir.stream import InputStreamError
    HACHOIR_AVAILABLE = True
except ImportError:
    HACHOIR_AVAILABLE = False

try:
    import mutagen
    from mutagen.mp3 import MP3
    from mutagen.flac import FLAC
    from mutagen.mp4 import MP4
    from mutagen.wave import WAVE
    from mutagen.oggvorbis import OggVorbis
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False

try:
    import pymediainfo
    MEDIAINFO_AVAILABLE = True
except ImportError:
    MEDIAINFO_AVAILABLE = False

class MetadataExtractor:
    """Base class for metadata extraction."""

    def extract(self, file_path):
        """
        Extract metadata from a file.

        Args:
            file_path (str): Path to the file

        Returns:
            dict: Dictionary containing metadata
        """
        raise NotImplementedError("Subclasses must implement extract method")


class ImageMetadataExtractor(MetadataExtractor):
    """Extract metadata from image files."""

    def extract(self, file_path):
        start_time = time.time()
        metadata = {
            "extracted_by": "ImageMetadataExtractor",
            "exif": {},
            "exif_detailed": {},
            "image_info": {}
        }

        try:
            # Method 1: Extract metadata using Pillow
            with Image.open(file_path) as img:
                # Basic image info
                metadata["image_info"] = {
                    "format": img.format,
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height,
                    "aspect_ratio": round(img.width / img.height, 2) if img.height != 0 else None,
                }

                # EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    if exif_data:
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)

                            # Convert bytes to string or ignore if not convertible
                            if isinstance(value, bytes):
                                try:
                                    value = value.decode('utf-8')
                                except UnicodeDecodeError:
                                    value = str(value)

                            metadata["exif"][tag] = value

            # Method 2: More detailed EXIF using exifread
            with open(file_path, 'rb') as f:
                exif_tags = exifread.process_file(f)

                # Organize EXIF data by category
                exif_by_category = {}
                gps_data = {}

                for tag, value in exif_tags.items():
                    # Extract category from tag name
                    if "GPS" in tag:
                        # Store GPS data separately
                        gps_data[tag] = str(value)
                    else:
                        # Get the category from the tag name
                        parts = tag.split()
                        category = parts[0] if len(parts) > 0 else "Other"

                        # Clean up category name
                        category = re.sub(r'[^\w\s]', '', category)

                        if category not in exif_by_category:
                            exif_by_category[category] = {}

                        # Store the tag value
                        tag_name = " ".join(parts[1:]) if len(parts) > 1 else tag
                        exif_by_category[category][tag_name] = str(value)

                # Add GPS coordinates if available
                if gps_data:
                    # Try to extract latitude and longitude
                    try:
                        lat_ref = gps_data.get("GPS GPSLatitudeRef", "N")
                        lat = self._convert_to_degrees(gps_data.get("GPS GPSLatitude", "0, 0, 0"))
                        if lat_ref == "S":
                            lat = -lat

                        lon_ref = gps_data.get("GPS GPSLongitudeRef", "E")
                        lon = self._convert_to_degrees(gps_data.get("GPS GPSLongitude", "0, 0, 0"))
                        if lon_ref == "W":
                            lon = -lon

                        metadata["gps_coordinates"] = {
                            "latitude": lat,
                            "longitude": lon,
                            "google_maps_url": f"https://www.google.com/maps/search/?api=1&query={lat},{lon}"
                        }
                    except Exception as e:
                        # GPS conversion failed
                        pass

                metadata["exif_detailed"] = exif_by_category

                # Extract any text found in the EXIF data
                text_content = []
                for category, tags in exif_by_category.items():
                    for tag, value in tags.items():
                        if isinstance(value, str) and len(value) > 10 and not value.startswith("["):
                            text_content.append(f"{tag}: {value}")

                if "ImageDescription" in metadata["exif"]:
                    text_content.append(f"Image Description: {metadata['exif']['ImageDescription']}")

                if "UserComment" in metadata["exif"]:
                    text_content.append(f"User Comment: {metadata['exif']['UserComment']}")

                if "Artist" in metadata["exif"]:
                    text_content.append(f"Artist: {metadata['exif']['Artist']}")

                if "Copyright" in metadata["exif"]:
                    text_content.append(f"Copyright: {metadata['exif']['Copyright']}")

                if text_content:
                    metadata["extracted_text"] = "\n".join(text_content)

            # Add steganography analysis if the image is a JPEG or PNG
            if metadata["image_info"]["format"] in ["JPEG", "PNG"]:
                try:
                    import cv2
                    import numpy as np
                    from scipy.fftpack import dct

                    # Check image size to avoid processing large images
                    img_size = metadata["image_info"]["width"] * metadata["image_info"]["height"]

                    # Always run the faster analysis methods
                    metadata["LSB_Analysis"] = self._lsb_analysis(file_path)
                    metadata["Chi_Square_Analysis"] = self._chi_square_analysis(file_path)

                    # Only run more intensive analysis for smaller images (under 1 million pixels)
                    if img_size < 1000000:  # 1000x1000 or equivalent
                        metadata["DCT_Analysis"] = self._dct_analysis(file_path)
                        metadata["Sample_Pair_Analysis"] = self._sample_pair_analysis(file_path)
                        metadata["RS_Analysis"] = self._rs_analysis(file_path)
                        metadata["PVD_Analysis"] = self._pvd_analysis(file_path)
                    else:
                        metadata["DCT_Analysis"] = "Skipped for large image"
                        metadata["Sample_Pair_Analysis"] = "Skipped for large image"
                        metadata["RS_Analysis"] = "Skipped for large image"
                        metadata["PVD_Analysis"] = "Skipped for large image"
                        metadata["Analysis_Note"] = "Some analyses were skipped due to large image size"
                except ImportError:
                    metadata["steganography_analysis"] = "Steganography analysis requires OpenCV and SciPy libraries"
                except Exception as e:
                    metadata["steganography_analysis_error"] = str(e)

        except UnidentifiedImageError:
            metadata["error"] = "Could not identify image file"
        except Exception as e:
            metadata["error"] = str(e)

        metadata["extraction_duration"] = int((time.time() - start_time) * 1000)
        return metadata

    def _convert_to_degrees(self, value):
        """Helper method to convert GPS coordinates from DMS format to decimal degrees."""
        if not value:
            return 0

        # Handle format like "[2, 29, 40.4453]"
        if isinstance(value, str):
            value = value.strip("[]")
            parts = [float(p.strip()) for p in value.split(",")]
        else:
            parts = value.values

        # Calculate degrees, minutes, seconds
        degrees = float(parts[0])
        minutes = float(parts[1]) / 60.0
        seconds = float(parts[2]) / 3600.0

        return degrees + minutes + seconds

    def _lsb_analysis(self, file_path):
        """Analyze image for LSB steganography."""
        try:
            import cv2
            import numpy as np

            img = cv2.imread(file_path)
            if img is None:
                return "Unable to read image file"

            lsb = img[:,:,0] % 2
            unusual_patterns = np.sum(lsb) / (img.shape[0] * img.shape[1])

            if unusual_patterns > 0.45 and unusual_patterns < 0.55:
                return "Suspicious: Possible LSB steganography detected"
            else:
                return "No obvious LSB steganography detected"
        except Exception as e:
            return f"Error in LSB analysis: {str(e)}"

    def _chi_square_analysis(self, file_path):
        """Analyze image using Chi-Square statistical test."""
        try:
            import cv2
            import numpy as np

            img = cv2.imread(file_path, 0)  # Read as grayscale
            if img is None:
                return "Unable to read image file"

            hist = cv2.calcHist([img], [0], None, [256], [0, 256])
            even_hist = hist[::2]
            odd_hist = hist[1::2]

            chi_square = np.sum((even_hist - odd_hist)**2 / (even_hist + odd_hist + 1e-6))

            if chi_square < 0.1:
                return "Suspicious: Possible steganography detected by Chi-Square analysis"
            else:
                return "No steganography detected by Chi-Square analysis"
        except Exception as e:
            return f"Error in Chi-Square analysis: {str(e)}"

    def _dct_analysis(self, file_path):
        """Analyze image using Discrete Cosine Transform."""
        try:
            import cv2
            import numpy as np
            from scipy.fftpack import dct

            img = cv2.imread(file_path, 0)  # Read as grayscale
            if img is None:
                return "Unable to read image file"

            img_dct = dct(dct(img.T, norm='ortho').T, norm='ortho')
            dct_values = np.abs(img_dct.flatten())

            threshold = np.percentile(dct_values, 99.95)
            suspicious_coeffs = np.sum(dct_values > threshold)

            if suspicious_coeffs > 100:
                return "Suspicious: Possible steganography detected by DCT analysis"
            else:
                return "No obvious steganography detected by DCT analysis"
        except Exception as e:
            return f"Error in DCT analysis: {str(e)}"

    def _sample_pair_analysis(self, file_path):
        """Analyze image using Sample Pair Analysis."""
        try:
            import cv2
            import numpy as np

            img = cv2.imread(file_path, 0)  # Read as grayscale
            if img is None:
                return "Unable to read image file"

            rows, cols = img.shape
            pairs = np.column_stack((img[:-1, :].flatten(), img[1:, :].flatten()))

            even_pairs = pairs[np.sum(pairs % 2, axis=1) == 0]
            odd_pairs = pairs[np.sum(pairs % 2, axis=1) == 1]

            beta = len(even_pairs) / (len(even_pairs) + len(odd_pairs))

            if 0.45 < beta < 0.55:
                return "Suspicious: Possible steganography detected by Sample Pair Analysis"
            else:
                return "No steganography detected by Sample Pair Analysis"
        except Exception as e:
            return f"Error in Sample Pair Analysis: {str(e)}"

    def _rs_analysis(self, file_path):
        """Analyze image using RS (Regular-Singular) Analysis with sampling for performance."""
        try:
            import cv2
            import numpy as np

            img = cv2.imread(file_path, 0)  # Read as grayscale
            if img is None:
                return "Unable to read image file"

            def flip_lsb(x):
                return x ^ 1

            rows, cols = img.shape

            # Use sampling for large images
            # Sample at most 2500 blocks (50x50 grid of 2x2 blocks)
            step_row = max(2, rows // 50)
            step_col = max(2, cols // 50)

            r_m, s_m, r_m_inv, s_m_inv = 0, 0, 0, 0

            for i in range(0, rows-1, step_row):
                if i + 2 > rows:
                    continue
                for j in range(0, cols-1, step_col):
                    if j + 2 > cols:
                        continue

                    block = img[i:i+2, j:j+2].astype(int)
                    flipped_block = np.vectorize(flip_lsb)(block)

                    diff = np.sum(np.abs(block - np.roll(block, 1, axis=1)))
                    diff_flipped = np.sum(np.abs(flipped_block - np.roll(flipped_block, 1, axis=1)))

                    if diff < diff_flipped:
                        r_m += 1
                    elif diff > diff_flipped:
                        s_m += 1

                    diff_inv = np.sum(np.abs(block - np.roll(block, -1, axis=1)))
                    diff_flipped_inv = np.sum(np.abs(flipped_block - np.roll(flipped_block, -1, axis=1)))

                    if diff_inv < diff_flipped_inv:
                        r_m_inv += 1
                    elif diff_inv > diff_flipped_inv:
                        s_m_inv += 1

            total = r_m + s_m
            total_inv = r_m_inv + s_m_inv

            if total == 0 or total_inv == 0:
                return "Unable to perform RS Analysis: insufficient data"

            d = abs((r_m - s_m) / total) - abs((r_m_inv - s_m_inv) / total_inv)

            if abs(d) < 0.05:
                return "Suspicious: Possible steganography detected by RS Analysis"
            else:
                return "No steganography detected by RS Analysis"
        except Exception as e:
            return f"Error in RS Analysis: {str(e)}"

    def _pvd_analysis(self, file_path):
        """Analyze image using Pixel Value Differencing with improved performance."""
        try:
            import cv2
            import numpy as np

            img = cv2.imread(file_path, 0)  # Read as grayscale
            if img is None:
                return "Unable to read image file"

            rows, cols = img.shape

            # For large images, sample only a portion
            if rows * cols > 1000000:  # For images larger than 1MP
                # Take a sample from the center of the image
                sample_size = 1000  # 1000x1000 sample
                start_row = max(0, (rows - sample_size) // 2)
                start_col = max(0, (cols - sample_size) // 2)
                end_row = min(rows, start_row + sample_size)
                end_col = min(cols, start_col + sample_size)
                img = img[start_row:end_row, start_col:end_col]
                rows, cols = img.shape

            # Vectorized implementation for horizontal differences
            horizontal_diffs = np.abs(img[:, 1:] - img[:, :-1])

            # Create and normalize histogram
            diff_hist = np.bincount(horizontal_diffs.flatten(), minlength=256)
            diff_hist = diff_hist / np.sum(diff_hist)

            # Check for unusual peaks in the histogram
            mean_val = np.mean(diff_hist)
            std_val = np.std(diff_hist)
            peaks = np.where(diff_hist > mean_val + 2 * std_val)[0]

            if len(peaks) > 5:  # Threshold can be adjusted
                return "Suspicious: Possible steganography detected by PVD Analysis"
            else:
                return "No obvious steganography detected by PVD Analysis"
        except Exception as e:
            return f"Error in PVD Analysis: {str(e)}"


class PDFMetadataExtractor(MetadataExtractor):
    """Extract metadata from PDF files."""

    def extract(self, file_path):
        start_time = time.time()
        metadata = {
            "extracted_by": "PDFMetadataExtractor",
            "pdf_info": {},
            "text_preview": ""
        }

        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                # Basic PDF info
                metadata["pdf_info"] = {
                    "num_pages": len(pdf_reader.pages),
                    "is_encrypted": pdf_reader.is_encrypted,
                }

                # Document info if available
                if pdf_reader.metadata:
                    for key, value in pdf_reader.metadata.items():
                        if key.startswith('/'):
                            key = key[1:]  # Remove leading slash
                        metadata["pdf_info"][key] = value

                # Extract full text content from all pages (up to a reasonable limit)
                full_text = []
                page_limit = min(10, len(pdf_reader.pages))  # Process up to 10 pages
                for i in range(page_limit):
                    try:
                        page_text = pdf_reader.pages[i].extract_text()
                        if page_text:
                            full_text.append(f"[Page {i+1}]\n{page_text}")
                    except Exception as e:
                        full_text.append(f"[Error extracting page {i+1}: {str(e)}]")

                # Join all text with page markers
                all_text = "\n\n".join(full_text)

                # Set both the preview (truncated) and the full extracted text
                metadata["text_preview"] = all_text[:5000] if all_text else ""  # 5000 char limit for preview
                metadata["extracted_text"] = all_text  # Store the complete text

        except Exception as e:
            metadata["error"] = str(e)

        metadata["extraction_time_ms"] = int((time.time() - start_time) * 1000)
        return metadata


class DocxMetadataExtractor(MetadataExtractor):
    """Extract metadata from DOCX files."""

    def extract(self, file_path):
        start_time = time.time()
        metadata = {
            "extracted_by": "DocxMetadataExtractor",
            "document_info": {},
            "text_preview": ""
        }

        try:
            doc = docx.Document(file_path)

            # Document properties
            core_properties = doc.core_properties
            metadata["document_info"] = {
                "author": core_properties.author,
                "created": str(core_properties.created) if core_properties.created else None,
                "last_modified_by": core_properties.last_modified_by,
                "modified": str(core_properties.modified) if core_properties.modified else None,
                "title": core_properties.title,
                "subject": core_properties.subject,
                "keywords": core_properties.keywords,
                "language": core_properties.language,
                "category": core_properties.category,
                "comments": core_properties.comments,
                "content_status": core_properties.content_status,
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections),
            }

            # Extract all text from the document
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    full_text.append(para.text)

            # Join all paragraphs with proper spacing
            all_text = "\n".join(full_text)

            # Set both the preview and full text
            metadata["text_preview"] = all_text[:5000] if all_text else ""  # 5000 char limit for preview
            metadata["extracted_text"] = all_text  # Store the complete text

        except Exception as e:
            metadata["error"] = str(e)

        metadata["extraction_time_ms"] = int((time.time() - start_time) * 1000)
        return metadata


class ExcelMetadataExtractor(MetadataExtractor):
    """Extract metadata from Excel files."""

    def extract(self, file_path):
        start_time = time.time()
        metadata = {
            "extracted_by": "ExcelMetadataExtractor",
            "workbook_info": {},
            "sheets": []
        }

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

            # Workbook properties
            metadata["workbook_info"] = {
                "sheet_names": wb.sheetnames,
                "sheet_count": len(wb.sheetnames),
                "properties": {
                    "creator": wb.properties.creator,
                    "last_modified_by": wb.properties.lastModifiedBy,
                    "created": str(wb.properties.created) if wb.properties.created else None,
                    "modified": str(wb.properties.modified) if wb.properties.modified else None,
                    "title": wb.properties.title,
                    "subject": wb.properties.subject,
                    "keywords": wb.properties.keywords,
                    "category": wb.properties.category,
                    "description": wb.properties.description
                }
            }

            # Get info for each sheet and extract text content
            full_text = []
            data_samples = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_info = {
                    "name": sheet_name,
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                }
                metadata["sheets"].append(sheet_info)

                # Extract a sample of data from each sheet (first 50 rows, first 10 columns)
                sheet_content = []
                sheet_content.append(f"[Sheet: {sheet_name}]")

                # Get header row if exists
                headers = []
                for j in range(1, min(11, sheet.max_column + 1)):
                    cell_value = sheet.cell(row=1, column=j).value
                    if cell_value is not None:
                        headers.append(str(cell_value))

                if headers:
                    sheet_content.append("Headers: " + " | ".join(headers))

                # Get a sample of data
                sample_rows = []
                for i in range(2, min(51, sheet.max_row + 1)):
                    row_data = []
                    for j in range(1, min(11, sheet.max_column + 1)):
                        cell_value = sheet.cell(row=i, column=j).value
                        if cell_value is not None:
                            row_data.append(str(cell_value))
                    if row_data:
                        sample_rows.append(" | ".join(row_data))

                if sample_rows:
                    sheet_content.append("Data Sample:")
                    sheet_content.extend(sample_rows[:20])  # Limit to 20 rows in the sample

                full_text.extend(sheet_content)

            # Join all sheet content with proper spacing
            all_text = "\n".join(full_text)

            # Set the extracted text
            metadata["text_preview"] = all_text[:5000] if all_text else ""  # 5000 char limit for preview
            metadata["extracted_text"] = all_text  # Store the complete text

        except Exception as e:
            metadata["error"] = str(e)

        metadata["extraction_time_ms"] = int((time.time() - start_time) * 1000)
        return metadata


class TextMetadataExtractor(MetadataExtractor):
    """Extract metadata from text files."""

    def extract(self, file_path):
        start_time = time.time()
        metadata = {
            "extracted_by": "TextMetadataExtractor",
            "text_info": {},
            "text_preview": ""
        }

        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                lines = content.splitlines()

                metadata["text_info"] = {
                    "line_count": len(lines),
                    "word_count": len(content.split()),
                    "character_count": len(content),
                    "file_size_bytes": os.path.getsize(file_path)
                }

                # Text preview (limited to 5000 chars)
                metadata["text_preview"] = content[:5000] if content else ""
                metadata["extracted_text"] = content  # Store the complete text

        except Exception as e:
            metadata["error"] = str(e)

        metadata["extraction_time_ms"] = int((time.time() - start_time) * 1000)
        return metadata


class AudioMetadataExtractor(MetadataExtractor):
    """Extract metadata from audio files using mutagen and hachoir."""

    def extract(self, file_path):
        start_time = time.time()
        metadata = {
            "extracted_by": "AudioMetadataExtractor",
            "audio_info": {},
            "tags": {},
            "technical_info": {}
        }

        # Get file size
        metadata["audio_info"]["file_size_bytes"] = os.path.getsize(file_path)

        # Extract metadata using mutagen if available
        if MUTAGEN_AVAILABLE:
            try:
                # Try to identify the audio format and use the appropriate class
                audio = None
                extension = os.path.splitext(file_path)[1].lower()

                if extension == '.mp3':
                    audio = MP3(file_path)
                elif extension == '.flac':
                    audio = FLAC(file_path)
                elif extension in ['.m4a', '.mp4', '.aac']:
                    audio = MP4(file_path)
                elif extension == '.wav':
                    audio = WAVE(file_path)
                elif extension in ['.ogg', '.oga']:
                    audio = OggVorbis(file_path)
                else:
                    # Generic audio file
                    audio = mutagen.File(file_path)

                if audio:
                    # Basic audio info
                    if hasattr(audio, 'info'):
                        info = audio.info
                        metadata["audio_info"].update({
                            "length_seconds": info.length,
                            "length_formatted": self._format_duration(info.length),
                            "bitrate": getattr(info, 'bitrate', None),
                            "sample_rate": getattr(info, 'sample_rate', None),
                            "channels": getattr(info, 'channels', None),
                            "codec": getattr(info, 'codec', None) or getattr(info, 'codec_name', None) or getattr(info, 'codec_description', None)
                        })

                    # Tags/metadata
                    if hasattr(audio, 'tags') and audio.tags:
                        for key, value in audio.tags.items():
                            # Handle different tag formats
                            if isinstance(value, list):
                                metadata["tags"][key] = value[0] if value else ""
                            else:
                                metadata["tags"][key] = str(value)
                    else:
                        # MP4 and some others have a different structure
                        for key, value in audio.items():
                            if isinstance(value, list):
                                metadata["tags"][key] = str(value[0]) if value else ""
                            else:
                                metadata["tags"][key] = str(value)
            except Exception as e:
                metadata["mutagen_error"] = str(e)
        else:
            metadata["mutagen_status"] = "Mutagen library not available"

        # Extract metadata using hachoir if available
        if HACHOIR_AVAILABLE:
            try:
                parser = hachoir.parser.createParser(file_path)
                if parser:
                    with parser:
                        metadata_extractor = hachoir.metadata.extractMetadata(parser)
                        if metadata_extractor:
                            for line in metadata_extractor.exportPlaintext():
                                if ': ' in line:
                                    key, value = line.split(': ', 1)
                                    # Clean up the key name
                                    key = key.strip().replace(' ', '_').lower()
                                    metadata["technical_info"][key] = value
            except Exception as e:
                metadata["hachoir_error"] = str(e)
        else:
            metadata["hachoir_status"] = "Hachoir library not available"

        # Extract metadata using pymediainfo if available
        if MEDIAINFO_AVAILABLE:
            try:
                media_info = pymediainfo.MediaInfo.parse(file_path)
                for track in media_info.tracks:
                    if track.track_type == 'Audio':
                        for key, value in track.__dict__.items():
                            if value and not key.startswith('_') and key not in ['track_type', 'track_id']:
                                # Format the key name
                                formatted_key = key.lower()
                                metadata["technical_info"][formatted_key] = value
            except Exception as e:
                metadata["mediainfo_error"] = str(e)
        else:
            metadata["mediainfo_status"] = "MediaInfo library not available"

        metadata["extraction_time_ms"] = int((time.time() - start_time) * 1000)
        return metadata

    def _format_duration(self, seconds):
        """Format duration in seconds to HH:MM:SS format."""
        minutes, seconds = divmod(int(seconds), 60)
        hours, minutes = divmod(minutes, 60)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


class VideoMetadataExtractor(MetadataExtractor):
    """Extract metadata from video files using pymediainfo and hachoir."""

    def extract(self, file_path):
        start_time = time.time()
        metadata = {
            "extracted_by": "VideoMetadataExtractor",
            "video_info": {},
            "audio_tracks": [],
            "video_tracks": [],
            "subtitle_tracks": [],
            "technical_info": {}
        }

        # Get file size
        metadata["video_info"]["file_size_bytes"] = os.path.getsize(file_path)

        # Extract metadata using pymediainfo if available
        if MEDIAINFO_AVAILABLE:
            try:
                media_info = pymediainfo.MediaInfo.parse(file_path)

                # Process general track
                for track in media_info.tracks:
                    if track.track_type == 'General':
                        metadata["video_info"].update({
                            "format": getattr(track, 'format', None),
                            "duration_ms": getattr(track, 'duration', None),
                            "duration_formatted": self._format_duration(getattr(track, 'duration', 0) / 1000 if hasattr(track, 'duration') else 0),
                            "overall_bit_rate": getattr(track, 'overall_bit_rate', None),
                            "encoded_date": getattr(track, 'encoded_date', None),
                            "writing_application": getattr(track, 'writing_application', None),
                            "writing_library": getattr(track, 'writing_library', None)
                        })

                        # Add all other general properties
                        for key, value in track.__dict__.items():
                            if value and not key.startswith('_') and key not in ['track_type', 'track_id']:
                                # Skip properties already added
                                if key not in ['format', 'duration', 'overall_bit_rate', 'encoded_date', 'writing_application', 'writing_library']:
                                    metadata["video_info"][key.lower()] = value

                    # Process video tracks
                    elif track.track_type == 'Video':
                        video_track = {
                            "format": getattr(track, 'format', None),
                            "codec": getattr(track, 'codec', None) or getattr(track, 'codec_id', None),
                            "width": getattr(track, 'width', None),
                            "height": getattr(track, 'height', None),
                            "aspect_ratio": getattr(track, 'display_aspect_ratio', None),
                            "frame_rate": getattr(track, 'frame_rate', None),
                            "bit_depth": getattr(track, 'bit_depth', None),
                            "scan_type": getattr(track, 'scan_type', None),
                            "color_space": getattr(track, 'color_space', None)
                        }

                        # Add all other video properties
                        for key, value in track.__dict__.items():
                            if value and not key.startswith('_') and key not in ['track_type', 'track_id']:
                                # Skip properties already added
                                if key not in ['format', 'codec', 'codec_id', 'width', 'height', 'display_aspect_ratio', 'frame_rate', 'bit_depth', 'scan_type', 'color_space']:
                                    video_track[key.lower()] = value

                        metadata["video_tracks"].append(video_track)

                    # Process audio tracks
                    elif track.track_type == 'Audio':
                        audio_track = {
                            "format": getattr(track, 'format', None),
                            "codec": getattr(track, 'codec', None) or getattr(track, 'codec_id', None),
                            "channels": getattr(track, 'channel_s', None),
                            "sample_rate": getattr(track, 'sampling_rate', None),
                            "bit_rate": getattr(track, 'bit_rate', None),
                            "language": getattr(track, 'language', None)
                        }

                        # Add all other audio properties
                        for key, value in track.__dict__.items():
                            if value and not key.startswith('_') and key not in ['track_type', 'track_id']:
                                # Skip properties already added
                                if key not in ['format', 'codec', 'codec_id', 'channel_s', 'sampling_rate', 'bit_rate', 'language']:
                                    audio_track[key.lower()] = value

                        metadata["audio_tracks"].append(audio_track)

                    # Process subtitle tracks
                    elif track.track_type == 'Text':
                        subtitle_track = {
                            "format": getattr(track, 'format', None),
                            "codec": getattr(track, 'codec', None) or getattr(track, 'codec_id', None),
                            "language": getattr(track, 'language', None)
                        }

                        # Add all other subtitle properties
                        for key, value in track.__dict__.items():
                            if value and not key.startswith('_') and key not in ['track_type', 'track_id']:
                                # Skip properties already added
                                if key not in ['format', 'codec', 'codec_id', 'language']:
                                    subtitle_track[key.lower()] = value

                        metadata["subtitle_tracks"].append(subtitle_track)
            except Exception as e:
                metadata["mediainfo_error"] = str(e)
        else:
            metadata["mediainfo_status"] = "MediaInfo library not available"

        # Extract metadata using hachoir if available
        if HACHOIR_AVAILABLE:
            try:
                parser = hachoir.parser.createParser(file_path)
                if parser:
                    with parser:
                        metadata_extractor = hachoir.metadata.extractMetadata(parser)
                        if metadata_extractor:
                            for line in metadata_extractor.exportPlaintext():
                                if ': ' in line:
                                    key, value = line.split(': ', 1)
                                    # Clean up the key name
                                    key = key.strip().replace(' ', '_').lower()
                                    metadata["technical_info"][key] = value
            except Exception as e:
                metadata["hachoir_error"] = str(e)
        else:
            metadata["hachoir_status"] = "Hachoir library not available"

        metadata["extraction_time_ms"] = int((time.time() - start_time) * 1000)
        return metadata

    def _format_duration(self, seconds):
        """Format duration in seconds to HH:MM:SS format."""
        minutes, seconds = divmod(int(seconds), 60)
        hours, minutes = divmod(minutes, 60)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def get_extractor_for_file(file_path):
    """
    Determine the appropriate extractor based on the file type.

    Args:
        file_path (str): Path to the file

    Returns:
        MetadataExtractor: An instance of the appropriate extractor class
    """
    mime = magic.Magic(mime=True)
    mime_type = mime.from_file(file_path)
    extension = os.path.splitext(file_path)[1].lower().lstrip('.')

    # Determine extractor based on MIME type and extension
    if mime_type.startswith('image/'):
        return ImageMetadataExtractor()
    elif mime_type == 'application/pdf' or extension == 'pdf':
        return PDFMetadataExtractor()
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or extension == 'docx':
        return DocxMetadataExtractor()
    elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                      'application/vnd.ms-excel'] or extension in ['xlsx', 'xls']:
        return ExcelMetadataExtractor()
    elif mime_type.startswith('text/') or extension in ['txt', 'csv', 'md', 'json', 'xml', 'html']:
        return TextMetadataExtractor()
    elif mime_type.startswith('audio/') or extension in ['mp3', 'wav', 'flac', 'ogg', 'oga', 'm4a', 'aac']:
        return AudioMetadataExtractor()
    elif mime_type.startswith('video/') or extension in ['mp4', 'avi', 'mov', 'mkv', 'webm', 'flv', 'wmv']:
        return VideoMetadataExtractor()
    else:
        # Default to a basic metadata extractor
        return MetadataExtractor()


def extract_metadata(file_path):
    """
    Extract metadata from a file using the appropriate extractor.

    Args:
        file_path (str): Path to the file

    Returns:
        dict: Dictionary containing metadata
    """
    extractor = get_extractor_for_file(file_path)
    try:
        metadata = extractor.extract(file_path)
        return metadata
    except NotImplementedError:
        # If using the base extractor, return basic file info
        return {
            "extracted_by": "BasicFileInfo",
            "file_info": {
                "file_size": os.path.getsize(file_path),
                "mime_type": magic.Magic(mime=True).from_file(file_path),
                "extension": os.path.splitext(file_path)[1].lower().lstrip('.'),
                "last_modified": time.ctime(os.path.getmtime(file_path)),
                "created": time.ctime(os.path.getctime(file_path))
            }
        }